import re
from docx import Document

def extract_volume_and_issue(docx_path):
    # 读取docx文件
    doc = Document(docx_path)

    # 定义正则表达式模式
    chinese_pattern = re.compile(r'\[\d+\] [^\u4e00-\u9fa5]')
    volume_issue_pattern = re.compile(r'(\d+\(\d+\))')

    # 存储匹配结果
    matches = []

    # 遍历每个段落
    for paragraph in doc.paragraphs:
        # 判断是否包含中文字符
        if re.search(chinese_pattern, paragraph.text):
            # 在包含中文字符的行中搜索匹配项
            current_matches = re.findall(volume_issue_pattern, paragraph.text)
            if current_matches:
                matches.extend(current_matches)

    return matches
                
def italicize_matches(results):
    doc = Document('1.docx')

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for result in results:
                if result in run.text:
                    # 检查匹配项是否在运行文本中
                    if result in run.text:
                        # 将匹配项所在的运行的文本设置为斜体
                        run.font.italic = True

    doc.save("modified_test.docx")




# 调用函数
results = extract_volume_and_issue("1.docx")

italicize_matches(results)
